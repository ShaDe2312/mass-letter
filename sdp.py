from tkinter import *
from tkinter import filedialog
from time import sleep
from openpyxl import load_workbook
import docx
import os
from os import mkdir
from docx2pdf import convert

def myclick():
    root.destroy()
    
def dummy():
    return


root=Tk()
root.title("Document automator")
root.geometry('300x300')
mainLabel=Label(root,text="Please select the template word file")
mainLabel.pack()
root.after(2500,dummy)

root.iconbitmap("C:\\Users\\rugve\\Downloads\\word.ico")
errorLabel=Label(root,text="Some error occurred. Please try this again")

root.filename=filedialog.askopenfilename(initialdir="C:/" , title = "Select word file" , filetypes=(("docx files", "*.docx"), ("doc files",".doc")))

if(len(root.filename)!=0):
    root.filename.replace('\\','\\\\')
    wordAddress=root.filename
    print(wordAddress)
    myLabel=Label(root,text="File input successfully\t.\n\nNow let us select the excel file\n")
    myLabel.pack()

else:
    errorLabel.pack()
    root.after(4000,root.destroy())
       
root.iconbitmap("C:\\Users\\rugve\\Downloads\\excel.ico")

root.filename=filedialog.askopenfilename(initialdir="C:/" , title = "Select excel file" , filetypes=(("excel files", "*.xlsx"), ("excel 97-03 files",".xlsm"), ("csv files",".csv")))

if(len(root.filename)!=0):
    root.filename.replace('\\','\\\\')
    excelAddress=root.filename
    print(excelAddress)
    myLabel=Label(root,text="File input successfully\t.\n\n")
    myLabel.pack()
    
else:
    errorLabel.pack()
    root.after(4000,root.destroy())
    
mybutton=Button(root,text="Press to Continue",padx=30,pady=10,fg="white",bg="#000000",command=myclick)
mybutton.pack()
root.mainloop()

#Taking data from excel file
#Using the openpyxl module
workbook= load_workbook(filename=excelAddress)
replaceKeys=[]
workbook.sheetnames

sheet=workbook.active
        
for values in sheet.iter_rows(min_row=2,values_only=True):
    replaceKeys.append(list(values))
    
    
finalIndex=excelAddress.rfind("/")

newDirPath=excelAddress[0:finalIndex+1] + "Output_Folder"
print(newDirPath)
if not os.path.exists(newDirPath):
    mkdir(newDirPath)
    
#Taking data from a word file

templateDoc=docx.Document("C:\\Users\\rugve\\Documents\\SDP TESTING\\Test_word.docx")
iterationDoc=[]

for i in range(len(replaceKeys)):
    iterationDoc.append(templateDoc)

for i in range(len(replaceKeys)):
    document=docx.Document(wordAddress)
    j=-1
    for paragraph in document.paragraphs:
        if '__' in paragraph.text:
            j+=1
            paragraph.text=paragraph.text.replace('__',replaceKeys[i][j])
    document.save(newDirPath + f"/Document {i+1}.docx")


    
